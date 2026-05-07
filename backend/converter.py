"""
Módulo de conversión de Word a HTML con referencias bidireccionales
"""

import subprocess
import shutil
import re
import os
from docx import Document
from bs4 import BeautifulSoup
from styles import CSS_REVISTAS, INTERACTIVE_JS


def process_article(input_file, index, save_folder, pandoc_folder, revista):
    """Procesa un artículo completo con referencias interactivas"""
    base_name = f"articulo_{index}"
    
    # Paso 1: Pandoc para ecuaciones e imágenes
    temp_html = os.path.join(pandoc_folder, f"{base_name}_temp.html")
    media_folder = os.path.join(pandoc_folder, f"{base_name}_media")
    
    subprocess.run([
        "pandoc", input_file, "-o", temp_html,
        "--extract-media", media_folder, "--mathml"
    ], check=True)
    
    # Paso 2: Crear HTML final
    final_html = os.path.join(save_folder, f"{base_name}.html")
    final_media_folder = os.path.join(save_folder, f"{base_name}_media")
    
    create_final_html(input_file, final_html, revista)
    
    # Paso 3: Copiar imágenes
    copy_media_files(media_folder, final_media_folder)
    
    # Paso 4: Insertar ecuaciones
    insert_equations(temp_html, final_html)
    
    # Paso 5: Actualizar rutas de imágenes
    update_image_paths(final_html, base_name, final_media_folder)


def create_final_html(input_file, output_file, revista):
    """Crea HTML con CSS embebido y referencias interactivas"""
    css_content = CSS_REVISTAS.get(revista, CSS_REVISTAS['ric'])
    
    soup = BeautifulSoup(
        f"""<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Artículo</title>
    <style>{css_content}</style>
</head>
<body>
{INTERACTIVE_JS}
</body>
</html>""",
        "html.parser"
    )
    
    body = soup.body
    doc = Document(input_file)
    tables = iter(doc.tables)
    
    paragraph_counter = 0
    in_references_section = False
    
    for para in doc.paragraphs:
        para_text = para.text.strip()
        if not para_text:
            continue
        
        paragraph_counter += 1
        para_id = f"para-{paragraph_counter}"
        
        # Detectar inicio de sección de referencias
        if re.match(r'^\s*REFERENCIAS\s*$', para_text, re.IGNORECASE):
            in_references_section = True
        
        p = determine_paragraph_style(para_text, soup, in_references_section)
        if p:
            p['id'] = para_id
            
            # Procesar referencias en el texto [1], [2], etc.
            if not in_references_section:
                process_inline_references(p, para_text, para_id, soup)
            else:
                # Si es una referencia, agregar ID único y botón de volver
                ref_match = re.match(r'^\[(\d+)\]', para_text)
                if ref_match:
                    ref_num = ref_match.group(1)
                    p['id'] = f"ref-{ref_num}"
                    p['class'].append('reference-item')
                    p.string = para_text
            
            body.append(p)
        
        # Insertar placeholder para imágenes
        if para_text.startswith("Figura"):
            img_tag = soup.new_tag('p', **{'class': 'revista_imagen-centrar'})
            img_name_tag = soup.new_tag('img', **{
                'class': 'img_responsive',
                'src': 'PLACEHOLDER_IMAGE',
                'alt': 'image.png'
            })
            img_tag.append(img_name_tag)
            p.insert_before(img_tag)
        
        # Insertar tablas
        if para_text.startswith("Tabla"):
            try:
                table = next(tables)
                table_tag = create_table(table, soup)
                p.insert_after(table_tag)
            except StopIteration:
                pass
    
    # Asignar clases especiales a títulos y autores
    assign_header_classes(soup)
    
    # Agregar DOI
    add_doi(soup)
    
    with open(output_file, 'w', encoding='utf-8') as file:
        file.write(soup.prettify())


def process_inline_references(p_tag, text, para_id, soup):
    """Convierte [1], [2] en enlaces interactivos"""
    ref_pattern = r'\[(\d+)\]'
    refs_found = re.findall(ref_pattern, text)
    
    if refs_found:
        p_tag.string = ''
        parts = re.split(ref_pattern, text)
        
        for i, part in enumerate(parts):
            if i % 2 == 0:
                # Texto normal
                if part:
                    p_tag.append(part)
            else:
                # Número de referencia
                ref_num = part
                ref_link = soup.new_tag('a', **{
                    'href': f'#ref-{ref_num}',
                    'class': 'ref-link',
                    'onclick': f"savePosition('{para_id}'); return true;"
                })
                ref_link.string = f'[{ref_num}]'
                p_tag.append(ref_link)
    else:
        p_tag.string = text


def determine_paragraph_style(para_text, soup, in_references=False):
    """Determina el estilo del párrafo según su contenido"""
    
    # Referencias
    if in_references and re.match(r'^\[\d+\]', para_text):
        return soup.new_tag('p', **{'class': 'revista_referencias'})
    
    # Títulos numerados
    if re.match(r'^\d+\.\s', para_text):
        return soup.new_tag('p', **{'class': 'revista_titulo1'})
    elif re.match(r'^\d+\.\d+\s', para_text):
        return soup.new_tag('p', **{'class': 'revista_titulo2'})
    elif re.match(r'^\d+\.\d+\.\d+\s', para_text):
        return soup.new_tag('p', **{'class': 'revista_titulo3'})
    elif re.match(r'^\d+\.\d+\.\d+\.\d+\s', para_text):
        return soup.new_tag('p', **{'class': 'revista_titulo4'})
    
    # Figuras y tablas
    elif re.match(r"Figura [0-9]+|Figura+", para_text):
        return soup.new_tag('p', **{'class': 'revista_piefototabla'})
    elif re.match(r"Fuente:+|Nota:", para_text):
        return soup.new_tag('p', **{'class': 'revista_piefototabla'})
    elif re.match(r"Tabla [0-9]+|Tabla+", para_text):
        return soup.new_tag('p', **{'class': 'revista_encabezadotabla'})
    
    # Secciones especiales
    elif re.search(r'\b(?:AGRADECIMIENTO|AGRADECIMIENTOS|CONFLICTO DE INTERESES|CONTRIBUCIÓN Y APROBACIÓN DE LOS AUTORES|ANEXOS|REFERENCIAS)\b', para_text):
        return soup.new_tag('p', **{'class': 'revista_titulo1'})
    
    # Ecuaciones
    elif re.match(r'\(\d+\)', para_text):
        return soup.new_tag('p', **{'class': 'insertar_ecuacion'})
    
    # Contenido normal
    else:
        return soup.new_tag('p', **{'class': 'revista_contenido'})


def create_table(table, soup):
    """Crea una tabla HTML desde tabla de Word"""
    table_tag = soup.new_tag('table', **{'border': '1', 'align': 'center', 'class': 'tabla'})
    tbody_tag = soup.new_tag('tbody', **{'align': 'center'})
    
    for row_index, row in enumerate(table.rows):
        row_tag = soup.new_tag('tr')
        if row_index == 0:
            row_tag['class'] = 'revista_encabezadotabla'
        
        for cell in row.cells:
            cell_tag = soup.new_tag('td')
            p_tag = soup.new_tag('p')
            p_tag.string = cell.text.strip()
            cell_tag.append(p_tag)
            row_tag.append(cell_tag)
        
        tbody_tag.append(row_tag)
    
    table_tag.append(tbody_tag)
    return table_tag


def assign_header_classes(soup):
    """Asigna clases a títulos y autores"""
    for idx, p_tag in enumerate(soup.find_all('p')):
        if idx == 0 or idx == 1:
            new_p_tag = soup.new_tag('p', **{'class': 'revista_tituloarticulo', 'id': p_tag.get('id', '')})
            new_p_tag.string = p_tag.string
            p_tag.replace_with(new_p_tag)
        elif idx == 2:
            new_p_tag = soup.new_tag('p', **{'class': 'revista_autores', 'id': p_tag.get('id', '')})
            new_p_tag.string = p_tag.string
            p_tag.replace_with(new_p_tag)


def add_doi(soup):
    """Agrega DOI antes del resumen"""
    for p_tag in soup.find_all('p'):
        if p_tag.string and p_tag.string.startswith(("RESUMEN", "Resumen")):
            doi_p = soup.new_tag('p', **{'class': 'revista_DOI'})
            doi_link = soup.new_tag('a', **{'href': 'https://doi.org/10.33412/rev-ric.v10.2.4040'})
            doi_link.string = "https://doi.org/10.33412/rev-ric.v10.2.4040 --> actualizar link del DOI"
            doi_p.append(doi_link)
            p_tag.insert_before(doi_p)
            break


def copy_media_files(source_folder, dest_folder):
    """Copia archivos multimedia"""
    if not os.path.exists(source_folder):
        return
    
    os.makedirs(dest_folder, exist_ok=True)
    media_source = os.path.join(source_folder, "media")
    
    if os.path.exists(media_source):
        for filename in os.listdir(media_source):
            source_file = os.path.join(media_source, filename)
            dest_file = os.path.join(dest_folder, filename)
            shutil.copy2(source_file, dest_file)


def insert_equations(temp_html_path, final_html_path):
    """Inserta ecuaciones del HTML de Pandoc"""
    with open(temp_html_path, 'r', encoding='utf-8') as f:
        temp_soup = BeautifulSoup(f.read(), 'html.parser')
    
    equations = temp_soup.find_all('math')
    if not equations:
        return
    
    with open(final_html_path, 'r', encoding='utf-8') as f:
        final_soup = BeautifulSoup(f.read(), 'html.parser')
    
def insert_equations(temp_html_path, final_html_path):
    """Inserta ecuaciones del HTML de Pandoc"""
    with open(temp_html_path, 'r', encoding='utf-8') as f:
        temp_soup = BeautifulSoup(f.read(), 'html.parser')
    
    equations = temp_soup.find_all('math')
    if not equations:
        return
    
    with open(final_html_path, 'r', encoding='utf-8') as f:
        final_soup = BeautifulSoup(f.read(), 'html.parser')
    
    equation_placeholders = final_soup.find_all('p', class_='insertar_ecuacion')
    
    for placeholder, equation in zip(equation_placeholders, equations):
        new_p = final_soup.new_tag('p')
        new_p.append(equation)
        placeholder.replace_with(new_p)
    
    with open(final_html_path, 'w', encoding='utf-8') as f:
        f.write(final_soup.prettify())


def update_image_paths(html_path, base_name, media_folder):
    """Actualiza las rutas de las imágenes"""
    with open(html_path, 'r', encoding='utf-8') as f:
        soup = BeautifulSoup(f.read(), 'html.parser')
    
    if not os.path.exists(media_folder):
        return
    
    image_files = sorted([f for f in os.listdir(media_folder) 
                         if f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.svg'))])
    
    img_tags = soup.find_all('img', src='PLACEHOLDER_IMAGE')
    
    for idx, img_tag in enumerate(img_tags):
        if idx < len(image_files):
            img_tag['src'] = f"{base_name}_media/{image_files[idx]}"
    
    with open(html_path, 'w', encoding='utf-8') as f:
        f.write(soup.prettify())